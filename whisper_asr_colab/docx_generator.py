#import logging
from pathlib import Path
from docx import Document

class DocxGenerator():
    def __init__(self, doc = None):
        # Create a new Word document
        template_path = Path(__file__).parent / 'templates' / 'diarized_transcription.docx'
        self.doc = doc if doc else Document(str(template_path))
        self.txtfilename = ""
        self.docfilename = ""
        self.styles = {
            "normal" : self.doc.styles['Normal'],
            "speaker" : self.doc.styles["Heading 1"],
            "ts1" : self.doc.styles['FirstParagraph'],
            "ts2" : self.doc.styles['SecondParagraph'],
        }

    def txt_to_word(self, txtfilename: str) -> None:
        self.txtfilename = txtfilename
        # Open the text file in read mode with utf8 encoding
        with open(txtfilename, 'r', encoding='utf8') as f:
            lines = f.readlines()

        # Process each line
        pline_count = 0
        line_break : bool = True
        for line in lines:
            # Blank line
            if line.strip() == "":
                continue
            # Time and speaker info
            if line.startswith('['):
                elements = line.split(' ')
                speaker = elements[3].strip()
                time = " ".join(elements[:3])
                self.doc.add_paragraph(speaker + ' ' + time, style=self.styles["speaker"])
                pline_count = 1
                continue
            # Transcript
            if pline_count == 1:
                self.doc.add_paragraph("○　", style=self.styles["ts1"])
            elif line_break:
                self.doc.add_paragraph("", style=self.styles["ts2"])
            text : str = line.strip()
            self.doc.paragraphs[-1].add_run(text)
            pline_count += 1
            # If the line is end with comma, do not break line
            line_break = ("、" != text[-1])

        self.docfilename = f"{Path(txtfilename).stem}.docx"
        self.doc.save(self.docfilename)


if __name__ == "__main__":
    import sys
    args = sys.argv
    for _file in args[1:]:
        print(f"Converting {_file} to docx...")
        doc = DocxGenerator()
        doc.txt_to_word(_file)
